#!/usr/bin/env python3
"""Build the downloadable Word version of the Cyber Essentials checklist.

Reads src/lib/content/cyber-essentials-checklist.json, the same file the
article renders from, so the document and the page cannot drift apart.

Requires python-docx, which is not a project dependency. Run it in a venv:

    python3 -m venv /tmp/docx && /tmp/docx/bin/pip install python-docx
    /tmp/docx/bin/python scripts/build-checklist-docx.py [OUTPUT_PATH]
"""

import json
import sys
from pathlib import Path

from docx import Document
from docx.enum.section import WD_ORIENT
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml import OxmlElement
from docx.oxml.ns import qn
from docx.shared import Cm, Pt, RGBColor

REPO = Path(__file__).resolve().parent.parent
DATA = REPO / "src" / "lib" / "content" / "cyber-essentials-checklist.json"
DEFAULT_OUT = REPO / "public" / "downloads" / "cyber-essentials-checklist.docx"

NAVY = RGBColor(0x0F, 0x20, 0x44)
EMERALD = RGBColor(0x05, 0x96, 0x69)
SLATE = RGBColor(0x47, 0x55, 0x69)
MUTED = RGBColor(0x64, 0x74, 0x8B)

COLUMNS = ["Task", "Evidence to record", "Owner", "Status"]
WIDTHS = [Cm(9.0), Cm(9.0), Cm(4.2), Cm(3.4)]


def shade(cell, hex_fill):
    el = OxmlElement("w:shd")
    el.set(qn("w:val"), "clear")
    el.set(qn("w:fill"), hex_fill)
    cell._tc.get_or_add_tcPr().append(el)


def run(paragraph, text, *, size=10, bold=False, color=SLATE, italic=False):
    r = paragraph.add_run(text)
    r.font.size = Pt(size)
    r.font.bold = bold
    r.font.italic = italic
    r.font.color.rgb = color
    r.font.name = "Calibri"
    return r


def spacer(doc, points=6):
    p = doc.add_paragraph()
    p.paragraph_format.space_before = Pt(0)
    p.paragraph_format.space_after = Pt(points)
    return p


def build(groups, out_path):
    doc = Document()

    section = doc.sections[0]
    section.orientation = WD_ORIENT.LANDSCAPE
    section.page_width, section.page_height = section.page_height, section.page_width
    for attr in ("top_margin", "bottom_margin", "left_margin", "right_margin"):
        setattr(section, attr, Cm(1.6))

    title = doc.add_paragraph()
    title.paragraph_format.space_after = Pt(2)
    run(title, "Cyber Essentials preparation checklist", size=18, bold=True, color=NAVY)

    sub = doc.add_paragraph()
    sub.paragraph_format.space_after = Pt(10)
    run(sub, "A preparation aid, not the certification application. ", size=9.5, color=MUTED)
    run(sub, "Work through scope first — every other answer depends on it.", size=9.5, color=MUTED, italic=True)

    fields = doc.add_paragraph()
    fields.paragraph_format.space_after = Pt(14)
    run(fields, "Organisation ", size=10, bold=True, color=NAVY)
    run(fields, "…" * 26 + "    ", size=10, color=MUTED)
    run(fields, "Completed by ", size=10, bold=True, color=NAVY)
    run(fields, "…" * 22 + "    ", size=10, color=MUTED)
    run(fields, "Date ", size=10, bold=True, color=NAVY)
    run(fields, "…" * 14, size=10, color=MUTED)

    total = 0

    for group in groups:
        heading = doc.add_paragraph()
        heading.paragraph_format.space_before = Pt(12)
        heading.paragraph_format.space_after = Pt(1)
        heading.paragraph_format.keep_with_next = True
        run(heading, group["stage"], size=12, bold=True, color=NAVY)

        purpose = doc.add_paragraph()
        purpose.paragraph_format.space_after = Pt(5)
        purpose.paragraph_format.keep_with_next = True
        run(purpose, group["purpose"], size=9, color=MUTED, italic=True)

        table = doc.add_table(rows=1, cols=len(COLUMNS))
        table.style = "Table Grid"
        table.alignment = WD_TABLE_ALIGNMENT.LEFT
        table.autofit = False

        header = table.rows[0]
        header.cells[0].paragraphs[0].paragraph_format.space_after = Pt(2)
        for idx, label in enumerate(COLUMNS):
            cell = header.cells[idx]
            cell.width = WIDTHS[idx]
            para = cell.paragraphs[0]
            para.paragraph_format.space_before = Pt(3)
            para.paragraph_format.space_after = Pt(3)
            run(para, label.upper(), size=8, bold=True, color=EMERALD)
            shade(cell, "F1F5F9")

        for item in group["items"]:
            total += 1
            row = table.add_row()
            values = [item["task"], item["evidence"], "", ""]
            for idx, value in enumerate(values):
                cell = row.cells[idx]
                cell.width = WIDTHS[idx]
                para = cell.paragraphs[0]
                para.paragraph_format.space_before = Pt(4)
                para.paragraph_format.space_after = Pt(4)
                if value:
                    run(para, value, size=9.5, color=NAVY if idx == 0 else SLATE)

    doc.add_page_break()

    closing = doc.add_paragraph()
    closing.paragraph_format.space_after = Pt(4)
    run(closing, f"{total} items across {len(groups)} stages.", size=10, bold=True, color=NAVY)

    note = doc.add_paragraph()
    note.paragraph_format.space_after = Pt(4)
    run(
        note,
        "Completing this checklist does not mean you will pass. It means you will know where you stand "
        "before you apply, rather than discovering gaps partway through an application.",
        size=9.5,
        color=SLATE,
    )

    disclaimer = doc.add_paragraph()
    disclaimer.paragraph_format.space_after = Pt(4)
    run(
        disclaimer,
        "BrightCert helps UK businesses prepare for Cyber Essentials by assessing readiness, identifying gaps "
        "and producing a practical report. BrightCert does not issue the official Cyber Essentials certificate. "
        "That comes from an IASME-licensed Certification Body.",
        size=8.5,
        color=MUTED,
    )

    source = doc.add_paragraph()
    source.alignment = WD_ALIGN_PARAGRAPH.LEFT
    run(source, "brightcert.co.uk/blog/cyber-essentials-checklist", size=8.5, bold=True, color=EMERALD)

    out_path.parent.mkdir(parents=True, exist_ok=True)
    doc.save(out_path)
    return total


if __name__ == "__main__":
    out = Path(sys.argv[1]) if len(sys.argv) > 1 else DEFAULT_OUT
    groups = json.loads(DATA.read_text())
    count = build(groups, out)
    print(f"wrote {out} — {count} items across {len(groups)} stages")
